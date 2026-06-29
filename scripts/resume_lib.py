# -*- coding: utf-8 -*-
"""sprite-skill resume + cover-letter renderer (python-docx).

This is the reusable engine. It is candidate-agnostic: all real content lives in a
separate `candidate.py` (copied from assets/candidate.example.py) so nothing personal
is baked into this public code. The structure encodes the career-advisor rules:

  Header -> Professional Summary -> Core Competencies -> Software (own section)
         -> Education (ABOVE experience) -> Professional Experience
         -> Volunteer & Cultural Engagement -> Additional Information

Hard rules enforced by the layout (see references/guardrails.md for the why):
  * Education sits above experience.
  * Software is its own section between Core Competencies and Education.
  * Each experience renders EXACTLY 3 bullets (volunteer is the deliberate exception: 2).
  * One page only -- verify with assemble.py before shipping.
  * drop_soft=True removes Volunteer + Additional for low-touch / non-people-facing
    roles, freeing the page for more detailed experience.
"""
import importlib.util
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

INK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0x0F, 0x4C, 0x81)
LINKC = "0F4C81"
FONT = "Calibri"


# ---------- low-level helpers ----------

def set_margins(doc, top=0.5, bottom=0.5, left=0.6, right=0.6):
    for s in doc.sections:
        s.top_margin = Inches(top); s.bottom_margin = Inches(bottom)
        s.left_margin = Inches(left); s.right_margin = Inches(right)


def base_style(doc, line_spacing=1.0):
    st = doc.styles["Normal"]
    st.font.name = FONT; st.font.size = Pt(10); st.font.color.rgb = INK
    pf = st.paragraph_format
    pf.space_after = Pt(0); pf.space_before = Pt(0); pf.line_spacing = line_spacing


def _run(p, text, size=10, bold=False, italic=False, color=INK, caps=False, spacing=None):
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = color
    rPr = r._element.get_or_add_rPr()
    if caps:
        c = OxmlElement("w:caps"); c.set(qn("w:val"), "true"); rPr.append(c)
    if spacing is not None:
        sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), str(spacing)); rPr.append(sp)
    return r


def add_hyperlink(p, text, url, size=9):
    part = p.part; r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), r_id)
    r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), LINKC); rPr.append(col)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size * 2)); rPr.append(sz)
    rf = OxmlElement("w:rFonts"); rf.set(qn("w:ascii"), FONT); rf.set(qn("w:hAnsi"), FONT); rPr.append(rf)
    r.append(rPr); t = OxmlElement("w:t"); t.text = text; r.append(t)
    h.append(r); p._p.append(h)


def bottom_border(p, color=LINKC, sz=8):
    pPr = p._p.get_or_add_pPr(); b = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "2"); bot.set(qn("w:color"), color)
    b.append(bot); pPr.append(b)


def section(doc, title, space_before=5):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(2)
    _run(p, title, size=11, bold=True, color=ACCENT, caps=True, spacing=24)
    bottom_border(p)
    return p


def entry(doc, role, org, dates, space_before=4):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(0)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(7.3), WD_TAB_ALIGNMENT.RIGHT)
    _run(p, role, size=10.5, bold=True)
    if org:
        _run(p, "   " + org, size=10.5, color=GREY)
    _run(p, "\t" + dates, size=9.5, italic=True, color=GREY)


def bullets(doc, items, space_after=1.0):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(space_after); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        _run(p, it, size=9.5)


def comp_line(doc, items):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for i, it in enumerate(items):
        if i:
            _run(p, "   |   ", size=9.5, color=ACCENT)
        _run(p, it, size=10)


# ---------- sections ----------

def header(doc, identity):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(1)
    _run(p, identity["name"], size=22, bold=True, color=INK, spacing=30)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after = Pt(4)
    _run(p2, identity["phone"] + "    |    ", size=9, color=GREY)
    add_hyperlink(p2, identity["email"], "mailto:" + identity["email"])
    for label, url in identity.get("links", []):
        _run(p2, "    |    ", size=9, color=GREY)
        add_hyperlink(p2, label, url)


def education(doc, education_list):
    section(doc, "Education")
    for name, deg, dates, course in education_list:
        entry(doc, name, "", dates, space_before=4)
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        _run(p, deg, size=10, italic=True, color=GREY)
        if course:
            pc = doc.add_paragraph(); pc.paragraph_format.space_after = Pt(0); pc.paragraph_format.left_indent = Inches(0.0)
            _run(pc, course, size=9.5, color=GREY)


def software_section(doc, software):
    section(doc, "Software")
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(0)
    _run(p, software, size=10)


def experience(doc, experiences, bullets_per_role, space_after=1.0):
    section(doc, "Professional Experience")
    for i, (role, org, dates) in enumerate(experiences):
        entry(doc, role, org, dates, space_before=(4 if i == 0 else 4))
        role_bullets = bullets_per_role[i] if i < len(bullets_per_role) else []
        bullets(doc, role_bullets, space_after=space_after)


def volunteer(doc, vol):
    if not vol:
        return
    role, org, dates, items = vol
    section(doc, "Volunteer & Cultural Engagement")
    entry(doc, role, org, dates, space_before=4)
    bullets(doc, items)


def additional(doc, brands, languages, hobbies):
    section(doc, "Additional Information")
    rows = []
    if brands:
        rows.append(("Brands", brands))
    if languages:
        rows.append(("Languages", languages))
    if hobbies:
        rows.append(("Hobbies", hobbies))
    for label, val in rows:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
        _run(p, label + ":  ", size=10, bold=True)
        _run(p, val, size=9.5, color=GREY)


# ---------- builders ----------

def build_resume(cand, variant, out, drop_soft=False):
    """Render one tailored resume .docx for `variant` from a loaded candidate module."""
    v = cand.VARIANTS[variant]
    line_spacing = v.get("line_spacing", 1.0)
    bullet_gap = v.get("bullet_space_after", 1.0)
    doc = Document(); base_style(doc, line_spacing=line_spacing); set_margins(doc)
    header(doc, cand.IDENTITY)
    section(doc, "Professional Summary")
    ps = doc.add_paragraph(); ps.paragraph_format.space_before = Pt(2); ps.paragraph_format.space_after = Pt(0)
    _run(ps, v["summary"], size=10)
    section(doc, "Core Competencies")
    comp_line(doc, v["competencies"])
    software_section(doc, v["software"])
    education(doc, cand.EDUCATION)
    experience(doc, cand.EXPERIENCES, v["bullets"], space_after=bullet_gap)
    if not drop_soft:
        volunteer(doc, getattr(cand, "VOLUNTEER", None))
        additional(doc, v.get("brands"), getattr(cand, "LANGUAGES", None), getattr(cand, "HOBBIES", None))
    doc.save(out)
    return out


def render_cover_letter(cand, md_path, out, line_spacing=1.12):
    """Render a one-page cover letter .docx: candidate letterhead + body read from `md_path`.

    The .md holds the (humanizer-polished) letter prose with [Company]/[Role] placeholders.
    Paragraphs are separated by blank lines; single newlines inside a block become line breaks
    (used for the sign-off).
    """
    import re
    import datetime
    text = open(md_path, encoding="utf-8").read().strip()
    blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]
    doc = Document(); base_style(doc, line_spacing=line_spacing)
    set_margins(doc, top=0.7, bottom=0.7, left=0.9, right=0.9)
    header(doc, cand.IDENTITY)
    today = datetime.date.today()
    date_str = "%d %s %d" % (today.day, today.strftime("%B"), today.year)
    pd = doc.add_paragraph(); pd.paragraph_format.space_before = Pt(6); pd.paragraph_format.space_after = Pt(8)
    _run(pd, date_str, size=10.5, color=GREY)
    for blk in blocks:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(9); p.paragraph_format.space_before = Pt(0)
        for j, ln in enumerate(blk.split("\n")):
            if j:
                p.add_run().add_break()
            _run(p, ln, size=10.5)
    doc.save(out)
    return out


def load_candidate(path):
    """Import a candidate.py from an arbitrary path and return the module."""
    path = os.path.abspath(path)
    spec = importlib.util.spec_from_file_location("candidate", path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod
